import docx
import re
from pathlib import Path

def md_to_docx(md_path: str, docx_path: str):
    doc = docx.Document()
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
        
    for line in content.split('\n'):
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('- **'):
            p = doc.add_paragraph(style='List Bullet')
            # Extract bold part
            match = re.match(r'- \*\*(.*?)\*\*(.*)', line)
            if match:
                bold_text = match.group(1)
                rest = match.group(2)
                p.add_run(bold_text).bold = True
                p.add_run(rest)
            else:
                p.add_run(line[2:])
        elif line.startswith('**'):
            p = doc.add_paragraph()
            p.add_run(line.strip('*')).bold = True
        elif line == '---':
            pass
        else:
            doc.add_paragraph(line)
            
    doc.save(docx_path)
    print("Saved docx to", docx_path)

md_file = r"C:\Users\harji\.gemini\antigravity\brain\e02b4e8d-8302-4b8a-826c-23460d8421e5\feature_list.md"
docx_file = r"C:\Users\harji\Downloads\KensaraAI_Feature_List.docx"
md_to_docx(md_file, docx_file)
